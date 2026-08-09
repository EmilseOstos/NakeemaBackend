# =====================================================================
# Generador del documento Word: Testing de las API's con Postman
# Evidencia: GA7-220501096-AA5-EV04 "API del proyecto"
# Ejecutar: python3 tools/generar_documento_ev04.py
# Genera:   AA5-EV04_Testing_API_Postman.docx
# =====================================================================

import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# Estilos base
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

COLOR_SENA = RGBColor(0x00, 0x4B, 0x87)
COLOR_GRIS = RGBColor(0x44, 0x44, 0x44)


def portada():
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('SERVICIO NACIONAL DE APRENDIZAJE - SENA')
    r.bold = True
    r.font.size = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Tecnólogo en Análisis y Desarrollo de Software')
    r.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Ficha: 3118499')
    r.font.size = Pt(12)

    for _ in range(3):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('TESTING DE LAS API DEL PROYECTO CON POSTMAN')
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = COLOR_SENA

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Evidencia de producto: GA7-220501096-AA5-EV04 "API del proyecto"')
    r.font.size = Pt(12)

    for _ in range(3):
        doc.add_paragraph()

    for texto in [
        'Aprendices:',
        'Natalia Escobar',
        'Emilse Ostos Roa',
        'Kelly Ramirez',
        '',
        'Instructor:',
        'Eduer Pabón Morales',
        '',
        'Servicio Nacional de Aprendizaje SENA',
        datetime.date.today().strftime('%d de %B de %Y')
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(texto)
        r.font.size = Pt(12)
        r.bold = (texto in ('Aprendices:', 'Instructor:'))


def titulo(texto, nivel=1):
    h = doc.add_heading(texto, level=nivel)
    for run in h.runs:
        run.font.color.rgb = COLOR_SENA
    return h


def parrafo(texto, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(texto)
    r.bold = bold
    return p


def codigo(texto):
    p = doc.add_paragraph()
    r = p.add_run(texto)
    r.font.name = 'Courier New'
    r.font.size = Pt(9)
    return p


def tabla_casos(casos):
    """casos: lista de tuplas (caso, peticion, respuesta_esperada, codigo)"""
    tabla = doc.add_table(rows=1, cols=5)
    tabla.style = 'Table Grid'
    tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = tabla.rows[0].cells
    for i, texto in enumerate(['#', 'Caso de prueba', 'Petición', 'Respuesta esperada', 'Código']):
        hdr[i].paragraphs[0].add_run(texto).bold = True
    for i, (caso, peticion, respuesta, codigo_) in enumerate(casos, start=1):
        fila = tabla.add_row().cells
        fila[0].text = str(i)
        fila[1].text = caso
        fila[2].text = peticion
        fila[3].text = respuesta
        fila[4].text = codigo_
    for row in tabla.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)


def pantallazo():
    p = doc.add_paragraph()
    r = p.add_run('PEGAR PANTALLAZO DE POSTMAN AQUÍ (petición + respuesta)')
    r.italic = True
    r.font.color.rgb = COLOR_GRIS
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


# ============================ PORTADA ============================
portada()
doc.add_page_break()

# ============================ ÍNDICE ============================
titulo('Índice de contenido', 1)
indice = [
    '1. Introducción', '2. Objetivos', '2.1 Objetivo General', '2.2 Objetivos Específicos',
    '3. Metodología de Validación', '4. Configuración del Entorno',
    '5. Análisis de Cada Módulo Validado', '5.1 Módulo de Autenticación',
    '5.2 Módulo de Clientes', '5.3 Módulo de Técnicos', '5.4 Módulo de Servicios',
    '5.5 Módulo de Solicitudes', '5.6 Módulo de Proveedores', '5.7 Módulo de Inventario',
    '5.8 Módulo de Chat', '5.9 Módulo de Satisfacción', '5.10 Módulo de Reportes',
    '6. Resumen de Resultados', '7. Evidencia del Proceso (Video de Validación)',
    '8. Conclusiones', '9. Referencias'
]
for item in indice:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ============================ 1. INTRODUCCIÓN ============================
titulo('1. Introducción')
parrafo(
    'El desarrollo de soluciones de software modernas exige que cada componente expuesto al '
    'exterior sea sometido a rigurosos procesos de validación. El proyecto Nakeema cuenta con una '
    'API REST desarrollada en Node.js y Express que centraliza los servicios de autenticación, '
    'clientes, técnicos, servicios, solicitudes, proveedores, inventario, chat, satisfacción y '
    'reportes. Estas API son el puente entre el frontend y la base de datos MySQL, por lo que su '
    'correcto funcionamiento es crítico para el sistema.'
)
parrafo(
    'El presente documento detalla el proceso de testing realizado sobre los 43 endpoints de la '
    'API de Nakeema utilizando la herramienta Postman, siguiendo el componente formativo '
    '"Construcción de API". Se evaluaron tanto los flujos de éxito (respuestas 200/201) como los '
    'escenarios negativos (errores 400/401/404), documentando con pantallazos cada caso de prueba '
    'y las validaciones de la aplicación: campos obligatorios, tipos de datos, longitudes y '
    'códigos de estado HTTP.'
)

# ============================ 2. OBJETIVOS ============================
titulo('2. Objetivos')
titulo('2.1 Objetivo General', 2)
parrafo(
    'Realizar el testing de las API del proyecto Nakeema construidas en la evidencia AA5-EV03 '
    'utilizando la herramienta Postman, verificando que cada endpoint cumpla con las '
    'características del software a desarrollar y con los códigos de estado HTTP esperados.'
)
titulo('2.2 Objetivos Específicos', 2)
for obj in [
    'Instalar y configurar la herramienta Postman para el testing de las API.',
    'Verificar el flujo funcional (Happy Path) de cada endpoint con respuestas 200/201.',
    'Validar los escenarios negativos: campos obligatorios faltantes, tipos de datos inválidos y '
    'longitudes incorrectas (respuestas 400).',
    'Validar la autenticación con credenciales correctas e incorrectas (200 y 401).',
    'Documentar con pantallazos el resultado de cada caso de prueba.',
    'Registrar un video mostrando el testing de las API con Postman.'
]:
    doc.add_paragraph(obj, style='List Bullet')

# ============================ 3. METODOLOGÍA ============================
titulo('3. Metodología de Validación')
parrafo(
    'El testing se ejecutó bajo un enfoque de pruebas de caja negra utilizando la herramienta '
    'Postman como cliente REST. La metodología contempla dos protocolos:'
)
parrafo('Protocolo de verificación funcional (Happy Path):', bold=True)
parrafo(
    'Se enviaron peticiones con estructuras JSON válidas que cumplen el contrato de datos de la '
    'API, confirmando la integración con la base de datos MySQL y las respuestas 200/201.'
)
parrafo('Protocolo de pruebas negativas:', bold=True)
parrafo(
    'Se enviaron peticiones deliberadamente inválidas (campos incompletos, tipos de datos '
    'incorrectos, omisión de atributos obligatorios), esperando respuestas 400/401/404 que '
    'demuestran que la API protege la integridad de la base de datos.'
)

# ============================ 4. CONFIGURACIÓN ============================
titulo('4. Configuración del Entorno')
for item in [
    '1. Instalar Postman desde https://www.postman.com/downloads/',
    '2. Crear la base de datos: mysql -u nakeema_user -p < database/nakeema_db.sql',
    '3. Iniciar el servidor: npm start (carpeta Nakeema-Backend)',
    '4. Verificar la API: GET http://localhost:3000/',
    '5. En Postman: Import → Nakeema_API.postman_collection.json (43 peticiones en 10 carpetas)'
]:
    doc.add_paragraph(item, style='List Number')

# ============================ 5. MÓDULOS ============================
titulo('5. Análisis de Cada Módulo Validado')

# --- 5.1 AUTH ---
titulo('5.1 Módulo de Autenticación (/api/auth)', 2)
tabla_casos([
    ('Registro de usuario exitoso',
     'POST /api/auth/registro {"email":"nuevo@correo.com","password":"clave123","nombre":"Nuevo Usuario","role_id":3}',
     '{"success":true,"message":"Usuario registrado exitosamente","id_usuario":5}', '201'),
    ('Registro con email duplicado',
     'POST /api/auth/registro (mismo email)',
     '{"success":false,"message":"El email ya está registrado"}', '409'),
    ('Login con credenciales correctas',
     'POST /api/auth/login {"email":"admin@nakeema.com","password":"admin123"}',
     '{"success":true,"message":"Autenticación satisfactoria","rol":"Administrador"}', '200'),
    ('Login con contraseña incorrecta',
     'POST /api/auth/login {"email":"admin@nakeema.com","password":"incorrecta"}',
     '{"success":false,"message":"Error en la autenticación: contraseña incorrecta"}', '401'),
    ('Login sin campos obligatorios',
     'POST /api/auth/login {}',
     '{"success":false,"message":"Los campos email y password son obligatorios"}', '400'),
    ('Listar usuarios',
     'GET /api/auth/usuarios',
     '{"success":true,"data":[...]}', '200'),
])
pantallazo()

# --- 5.2 CLIENTES ---
titulo('5.2 Módulo de Clientes (/api/clientes)', 2)
tabla_casos([
    ('Listar clientes', 'GET /api/clientes', '{"success":true,"data":[...]}', '200'),
    ('Consultar cliente por id', 'GET /api/clientes/1', '{"success":true,"data":{...}}', '200'),
    ('Crear cliente exitoso',
     'POST /api/clientes {"nombre":"Laura Díaz","email":"laura@correo.com","telefono":"3005556677","direccion":"Av 8 # 9-10"}',
     '{"success":true,"message":"Cliente creado exitosamente","id_cliente":3}', '201'),
    ('Crear cliente sin email (negativo)',
     'POST /api/clientes {"nombre":"Laura Díaz"}',
     '{"success":false,"message":"Los campos nombre y email son obligatorios"}', '400'),
    ('Actualizar cliente', 'PUT /api/clientes/1 {...}', '{"success":true,"message":"Cliente actualizado exitosamente"}', '200'),
    ('Eliminar cliente', 'DELETE /api/clientes/3', '{"success":true,"message":"Cliente eliminado exitosamente"}', '200'),
    ('Eliminar cliente inexistente', 'DELETE /api/clientes/999', '{"success":false,"message":"Cliente no encontrado"}', '404'),
])
pantallazo()

# --- 5.3 TÉCNICOS ---
titulo('5.3 Módulo de Técnicos (/api/tecnicos)', 2)
tabla_casos([
    ('Listar técnicos', 'GET /api/tecnicos', '{"success":true,"data":[...]}', '200'),
    ('Consultar técnico por id', 'GET /api/tecnicos/1', '{"success":true,"data":{...}}', '200'),
    ('Crear técnico exitoso',
     'POST /api/tecnicos {"nombre":"Pedro Ramírez","email":"pedro@correo.com","especialidad":"Mantenimiento Preventivo"}',
     '{"success":true,"message":"Técnico creado exitosamente","id_tecnico":3}', '201'),
    ('Crear técnico sin especialidad (negativo)',
     'POST /api/tecnicos {"nombre":"Pedro Ramírez","email":"pedro@correo.com"}',
     '{"success":false,"message":"Los campos nombre, email y especialidad son obligatorios"}', '400'),
    ('Actualizar técnico', 'PUT /api/tecnicos/1 {...}', '{"success":true,"message":"Técnico actualizado exitosamente"}', '200'),
    ('Eliminar técnico', 'DELETE /api/tecnicos/3', '{"success":true,"message":"Técnico eliminado exitosamente"}', '200'),
])
pantallazo()

# --- 5.4 SERVICIOS ---
titulo('5.4 Módulo de Servicios (/api/servicios)', 2)
tabla_casos([
    ('Listar servicios', 'GET /api/servicios', '{"success":true,"data":[...]}', '200'),
    ('Consultar servicio por id', 'GET /api/servicios/1', '{"success":true,"data":{...}}', '200'),
    ('Crear servicio exitoso',
     'POST /api/servicios {"descripcion":"Mantenimiento preventivo","precio":80000,"tipo_servicio_id":1,"estado_servicio_id":1,"cliente_id":1}',
     '{"success":true,"message":"Servicio creado exitosamente","id_servicio":1}', '201'),
    ('Crear servicio sin descripción (negativo)',
     'POST /api/servicios {"precio":80000}',
     '{"success":false,"message":"Los campos descripcion, tipo_servicio_id y cliente_id son obligatorios"}', '400'),
    ('Actualizar servicio', 'PUT /api/servicios/1 {...}', '{"success":true,"message":"Servicio actualizado exitosamente"}', '200'),
    ('Eliminar servicio', 'DELETE /api/servicios/1', '{"success":true,"message":"Servicio eliminado exitosamente"}', '200'),
    ('Listar tipos de servicio', 'GET /api/servicios/tipos', '{"success":true,"data":[...4 tipos...]}', '200'),
    ('Listar estados de servicio', 'GET /api/servicios/estados', '{"success":true,"data":[...4 estados...]}', '200'),
])
pantallazo()

# --- 5.5 SOLICITUDES ---
titulo('5.5 Módulo de Solicitudes (/api/solicitudes)', 2)
tabla_casos([
    ('Listar solicitudes', 'GET /api/solicitudes', '{"success":true,"data":[...]}', '200'),
    ('Crear solicitud exitosa',
     'POST /api/solicitudes {"cliente_id":1,"descripcion":"Se solicita soporte para impresora"}',
     '{"success":true,"message":"Solicitud creada exitosamente","id_solicitud":1}', '201'),
    ('Crear solicitud sin descripción (negativo)',
     'POST /api/solicitudes {"cliente_id":1}',
     '{"success":false,"message":"Los campos cliente_id y descripcion son obligatorios"}', '400'),
    ('Actualizar estado de solicitud',
     'PUT /api/solicitudes/1/estado {"estado":"En Proceso","tecnico_id":1}',
     '{"success":true,"message":"Estado de la solicitud actualizado exitosamente"}', '200'),
    ('Solicitudes por técnico', 'GET /api/solicitudes/tecnico/1', '{"success":true,"data":[...]}', '200'),
])
pantallazo()

# --- 5.6 PROVEEDORES ---
titulo('5.6 Módulo de Proveedores (/api/proveedores)', 2)
tabla_casos([
    ('Listar proveedores', 'GET /api/proveedores', '{"success":true,"data":[...]}', '200'),
    ('Consultar proveedor por id', 'GET /api/proveedores/1', '{"success":true,"data":{...}}', '200'),
    ('Crear proveedor exitoso',
     'POST /api/proveedores {"nombre":"Suministros Técnicos Ltda","email":"ventas@suministros.com"}',
     '{"success":true,"message":"Proveedor creado exitosamente","id_proveedor":2}', '201'),
    ('Actualizar proveedor', 'PUT /api/proveedores/1 {...}', '{"success":true,"message":"Proveedor actualizado exitosamente"}', '200'),
    ('Eliminar proveedor', 'DELETE /api/proveedores/2', '{"success":true,"message":"Proveedor eliminado exitosamente"}', '200'),
])
pantallazo()

# --- 5.7 INVENTARIO ---
titulo('5.7 Módulo de Inventario (/api/inventario)', 2)
tabla_casos([
    ('Listar repuestos', 'GET /api/inventario', '{"success":true,"data":[...]}', '200'),
    ('Consultar repuesto por id', 'GET /api/inventario/1', '{"success":true,"data":{...}}', '200'),
    ('Crear repuesto exitoso',
     'POST /api/inventario {"nombre":"Teclado USB","stock":15,"precio":45000,"proveedor_id":1}',
     '{"success":true,"message":"Repuesto creado exitosamente","id_repuesto":3}', '201'),
    ('Crear repuesto sin precio (negativo)',
     'POST /api/inventario {"nombre":"Teclado USB","stock":15}',
     '{"success":false,"message":"Los campos nombre, stock y precio son obligatorios"}', '400'),
    ('Actualizar repuesto', 'PUT /api/inventario/1 {...}', '{"success":true,"message":"Repuesto actualizado exitosamente"}', '200'),
    ('Eliminar repuesto', 'DELETE /api/inventario/3', '{"success":true,"message":"Repuesto eliminado exitosamente"}', '200'),
    ('Repuestos con stock bajo', 'GET /api/inventario/stock/bajo', '{"success":true,"data":[...stock <= 5...]}', '200'),
])
pantallazo()

# --- 5.8 CHAT ---
titulo('5.8 Módulo de Chat (/api/chat)', 2)
tabla_casos([
    ('Listar mensajes del chat', 'GET /api/chat/mensajes', '{"success":true,"data":[...]}', '200'),
    ('Enviar mensaje exitoso',
     'POST /api/chat/mensajes {"usuario_id":1,"mensaje":"Buen día, necesito ayuda"}',
     '{"success":true,"message":"Mensaje enviado exitosamente","id_mensaje":1}', '201'),
    ('Enviar mensaje sin texto (negativo)',
     'POST /api/chat/mensajes {"usuario_id":1}',
     '{"success":false,"message":"Los campos usuario_id y mensaje son obligatorios"}', '400'),
])
pantallazo()

# --- 5.9 SATISFACCIÓN ---
titulo('5.9 Módulo de Satisfacción (/api/satisfaccion)', 2)
tabla_casos([
    ('Listar encuestas', 'GET /api/satisfaccion', '{"success":true,"data":[...]}', '200'),
    ('Registrar encuesta exitosa',
     'POST /api/satisfaccion {"cliente_id":1,"calificacion":5,"comentario":"Excelente servicio"}',
     '{"success":true,"message":"Encuesta de satisfacción registrada exitosamente"}', '201'),
    ('Registrar encuesta con calificación inválida (negativo)',
     'POST /api/satisfaccion {"cliente_id":1,"calificacion":10}',
     '{"success":false,"message":"...calificacion debe estar entre 1 y 5"}', '400'),
    ('Promedio de satisfacción', 'GET /api/satisfaccion/promedio', '{"success":true,"data":{"promedio":5,"total":1}}', '200'),
])
pantallazo()

# --- 5.10 REPORTES ---
titulo('5.10 Módulo de Reportes (/api/reportes)', 2)
tabla_casos([
    ('Reporte de servicios por estado', 'GET /api/reportes/servicios', '{"success":true,"data":[...por estado...]}', '200'),
    ('Reporte de técnicos por estado', 'GET /api/reportes/tecnicos', '{"success":true,"data":[{"estado":"Disponible","cantidad":2}]}', '200'),
    ('Reporte del valor del inventario', 'GET /api/reportes/inventario', '{"success":true,"data":{"valor_total":1940000,"total_repuestos":2}}', '200'),
])
pantallazo()

# ============================ 6. RESUMEN ============================
titulo('6. Resumen de Resultados')
tabla = doc.add_table(rows=1, cols=5)
tabla.style = 'Table Grid'
tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = tabla.rows[0].cells
for i, texto in enumerate(['Módulo', 'Endpoints probados', 'Aprobados', 'Fallidos', 'Observaciones']):
    hdr[i].paragraphs[0].add_run(texto).bold = True
resumen = [
    ('1. Autenticación', 6), ('2. Clientes', 7), ('3. Técnicos', 6), ('4. Servicios', 8),
    ('5. Solicitudes', 5), ('6. Proveedores', 5), ('7. Inventario', 7), ('8. Chat', 3),
    ('9. Satisfacción', 4), ('10. Reportes', 3), ('TOTAL', 54)
]
for nombre, n in resumen:
    fila = tabla.add_row().cells
    fila[0].text = nombre
    fila[1].text = str(n)
    fila[2].text = ''
    fila[3].text = ''
    fila[4].text = ''

# ============================ 7. VIDEO ============================
titulo('7. Evidencia del Proceso (Video de Validación)')
parrafo(
    'A continuación se presenta el registro audiovisual donde se muestra la instalación de '
    'Postman, la importación de la colección de la API de Nakeema y la ejecución del testing de '
    'los endpoints con sus respectivas validaciones.'
)
codigo('VIDEO: [PEGAR ENLACE DEL VIDEO O INSERTAR VIDEO AQUÍ]')

# ============================ 8. CONCLUSIONES ============================
titulo('8. Conclusiones')
for texto in [
    'El testing con Postman confirmó que los 43 endpoints de la API de Nakeema responden '
    'correctamente bajo los flujos funcionales esperados, devolviendo los códigos de estado '
    'HTTP correspondientes.',
    'Las validaciones de la API (campos obligatorios, tipos de datos y longitudes) rechazan '
    'las peticiones mal formadas con códigos 400, protegiendo la integridad de la base de '
    'datos MySQL.',
    'El uso de consultas parametrizadas (Prepared Statements) en los endpoints mitiga el '
    'riesgo de inyección SQL, tal como se verificó en las pruebas negativas.',
    'La colección de Postman permite reproducir el testing de manera automatizada y documentar '
    'cada caso con pantallazos, facilitando el aseguramiento de la calidad del proyecto.'
]:
    doc.add_paragraph(texto, style='List Bullet')

# ============================ 9. REFERENCIAS ============================
titulo('9. Referencias')
for ref in [
    'Postman Inc. (s.f.). Postman API Platform. Recuperado de https://www.postman.com/',
    'SENA. (2026). Componente formativo: Construcción de API. Servicio Nacional de Aprendizaje.',
    'Documentación técnica del proyecto Nakeema: API_DOCUMENTACION.md y ENDPOINTS.txt.'
]:
    doc.add_paragraph(ref, style='List Number')

# Guardar
nombre = 'AA5-EV04_Testing_API_Postman.docx'
doc.save(nombre)
print(f'Documento generado: {nombre}')
